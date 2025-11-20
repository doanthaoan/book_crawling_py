import requests
from bs4 import BeautifulSoup
import time
import random
from colorama import Fore, Back, Style, init
import os
import html
from docx import Document
import re

init(autoreset=True)

BOOK_PATH = "truyen"
LOGS_PATH = "logs"
BOOK_DOMAIN = 'https://dichngay.com'

if not os.path.exists(BOOK_PATH):
    os.makedirs(BOOK_PATH)
if not os.path.exists(LOGS_PATH):
    os.makedirs(LOGS_PATH)

# Get chapter list from localfile
file_name = input("Đường dẫn file mục lục: ")
book_name = input("Tên truyện: ")
output_name_html = f"{book_name}.html"
output_name_docx = f"{book_name}.docx"
success_log_name = f"{book_name}_success.txt"
failure_log_name = f"{book_name}_failure.txt"
output_path_html = os.path.join(BOOK_PATH, output_name_html)
output_path_docx = os.path.join(BOOK_PATH, output_name_docx)
success_log_path = os.path.join(LOGS_PATH, success_log_name)
failure_log_path = os.path.join(LOGS_PATH, failure_log_name)
file_path = os.path.join('book_source', file_name)

with open(file_path, 'r', encoding='utf-8') as file:
    html_content = file.read()
soup = BeautifulSoup(html_content, 'html.parser')

chapter_tags = soup.select('li.mulu > a')

url_list = [a['href'] for a in chapter_tags if a.has_attr('href')]
# url_list = [
#     "https://dichngay.com/translate?u=https%3A%2F%2Fwww.52shuku.vip%2Fjiakong%2Fb%2FbjP0W_3.html&bid=Y5xr5FS4CBSVnOB7&un="
# ]
# print(url_list)
total_time = 0
# Extract content & save to html file
def clean_text(text):
    # Xóa NULL bytes (\x00) và các ký tự điều khiển không hợp lệ trong XML
    cleaned_text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    return cleaned_text

def append_to_docx(docx_filename, chapter_title, chapter_content):
    # Kiểm tra xem file đã tồn tại chưa
    if os.path.exists(docx_filename):
        doc = Document(docx_filename)  # Mở file nếu tồn tại
    else:
        doc = Document()  # Tạo tài liệu mới nếu chưa có file
    
    doc.add_heading(clean_text(chapter_title.get_text()), level=1)
    
    # Thêm nội dung mới vào tài liệu
    for paragraph in chapter_content:
        doc.add_paragraph(clean_text(paragraph.get_text()))
    
    # Lưu lại tài liệu ngay lập tức
    try:    
        doc.save(docx_filename)
        print(f"{Fore.GREEN}Lưu thành công DOCX {Fore.YELLOW}{chapter_title.get_text()}")
    except PermissionError:
        print(f"{Fore.RED}Lỗi: Không thể ghi file '{docx_filename}' do quyền hạn hoặc file đang được sử dụng.")
    except FileNotFoundError:
        print(f"{Fore.RED}Lỗi: Thư mục đích không tồn tại. Kiểm tra lại đường dẫn lưu file.")
    except Exception as e:
        print(f"{Fore.RED}Đã có lỗi xảy ra khi lưu file: {e}")

# Mark the new success log
with open(success_log_path, 'a') as success_log:
    success_log.write(f"\n=====================\n{time.strftime('%X %x')}: START NEW\n=====================\n")
# Mark the new fail log
with open(failure_log_path, 'a') as failure_log:
    failure_log.write(f"\n=====================\n{time.strftime('%X %x')}: START NEW\n=====================\n")

success_counter = 0
fail_counter = 0    

with open(output_path_html,'w', encoding='utf-8') as output_files:
    output_files.write('<html><body>\n')
    
    for index, url in enumerate(url_list):
        start_time = time.time()
        if "http" not in url:
            url = BOOK_DOMAIN + url
        try:
            print(f"=== Bắt đầu tải chương: \n{Fore.YELLOW}{url} ===")
            response = requests.get(url)
            end_request =  time.time()
            request_time = end_request - start_time
            print(f"Request time: {request_time:.2f} s")
            # print(response.text)
            if response.status_code == 200:
                html_content = response.text
                soup = BeautifulSoup(html_content, 'html.parser')
                
                iframe = soup.find('iframe')
                # decoded_content = html.unescape(iframe)
                inner_soup = BeautifulSoup(iframe["srcdoc"], 'html.parser')
                
                # print(inner_soup.find('article', {'class','article-content'}))
                chapter_title = inner_soup.find('h1', {'class', 'article-title'})
                chapter_content = inner_soup.find('article', {'class','article-content'})
               
                # HTML
                output_files.write(f"{str(chapter_title)}\n")
                output_files.write(str(chapter_content))
                # DOCX
                docx_content = chapter_content.find_all('p')
                [[a.replace_with(" " + a.get_text()) for a in p.find_all('a')] for p in docx_content]

                append_to_docx(output_path_docx, chapter_title, docx_content)
                
                # doc.add_heading(clean_text(chapter_title.get_text()), level=1)
                
                # for paragraph in docx_content:
                #     doc.add_paragraph(clean_text(paragraph.get_text()))
                
                success_counter += 1
                print(f"{Fore.GREEN}Lưu thành công HTML {Fore.YELLOW}{chapter_title.get_text()}: {Fore.WHITE}{url}")
                print(f"{Fore.WHITE}Tiến độ: {Fore.GREEN}{index+1}/{len(url_list)}")
                with open(success_log_path, 'a') as success_log:
                    success_log.write(f"{time.strftime('%X %x')}: {url}\n")
            else:
                print(f"{Fore.RED}Không thể truy cập {Fore.WHITE}{url}. Mã lỗi: {response.status_code}")
                with open(failure_log_path, 'a') as failure_log:
                    failure_log.write(f"{url}\n")
        except requests.exceptions.RequestException as e:
            print(f"{Fore.RED}Lỗi khi truy cập {Fore.WHITE}{url}: {e}")
            with open(failure_log_path, 'a') as failure_log:
                    failure_log.write(f"{url}\n")
        
        time.sleep(random.randint(4,6))
        end_time = time.time()
        duration = end_time - start_time
        total_time += duration
        print(f"Thời gian xử lý: {duration:.2f} s\n")
    
    output_files.write('</body></html>')
    
print(f"Tổng thời gian: {Fore.GREEN}{total_time:.2f} s")