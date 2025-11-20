import json
import os
import re
import requests
from requests.cookies import RequestsCookieJar
import time
import random
from docx import Document
from colorama import Fore, Back, Style, init
init(autoreset=True)

# API_ENDPOINT = "https://dichtienghoa.com/transtext"
# API_ENDPOINT = "https://api.dichnhanh.com/"
API_ENDPOINT = "https://dichngay.com/translate/text"
def append_to_docx(docx_filename, chapter_content):
    # Kiểm tra xem file đã tồn tại chưa
    if os.path.exists(docx_filename):
        doc = Document(docx_filename)  # Mở file nếu tồn tại
    else:
        doc = Document()  # Tạo tài liệu mới nếu chưa có file
    
    # doc.add_heading(clean_text(chapter_title.get_text()), level=1)
    
    # Thêm nội dung mới vào tài liệu
    lines = chapter_content.split('\n', 1)
    title = lines[0]  # The first line is the title
    content = lines[1] if len(lines) > 1 else ""
    doc.add_heading(title, level=1)
    # for paragraph in chapter_content:
    for paragraph in content.split('\n'):
        doc.add_paragraph((paragraph))
    
    # Lưu lại tài liệu ngay lập tức
    try:    
        doc.save(docx_filename)
        print(f"{Fore.GREEN}Lưu thành công DOCX {Fore.YELLOW}{title}")
    except PermissionError:
        print(f"{Fore.RED}Lỗi: Không thể ghi file '{docx_filename}' do quyền hạn hoặc file đang được sử dụng.")
    except FileNotFoundError:
        print(f"{Fore.RED}Lỗi: Thư mục đích không tồn tại. Kiểm tra lại đường dẫn lưu file.")
    except Exception as e:
        print(f"{Fore.RED}Đã có lỗi xảy ra khi lưu file: {e}")
        
# file_content = open("D:\\末世大佬穿到古代被流放.txt","r",  encoding='utf-8').read()
file_content = open("trieuhoansu.txt","r",  encoding='utf-8').read()
# Use regular expression to split the file into chapters
# Match pattern: chapter numbers at the start of a line
chapters = re.split(r"(?=\b\d+\.\s第\d+章)", file_content)
# Chapter with start with "第" and end with "章"
# chapters = re.split(r"(?=第\d+章)", file_content)

# Remove any empty strings and strip extra whitespace
chapters = [chapter.strip() for chapter in chapters if chapter.strip()]

# print(chapters[1])

# Tạo cookie jar để lưu trữ cookies
jar = RequestsCookieJar()
jar.set("cf_clearance", "wikv3C42Zu8jaA.d_VD06RJuRDOiILEucaWprnLy3Cs-1758984402-1.2.1.1-6m_azqQeYXtvSkGUnzug4JyAc7YP.lnOnuVoUIb27ZIRL_d7AD45SOHro7_s6fwrfRgQVvP_bMsIRasHAxkk7KlNsjp_VG0tvpxUT7ZkziuxJg5YJ.5nA3Xc0dFdMV6Bl.YLNdQVDRezyKhSNf4NeHNNedd_kavkcNqbulatsTu3LvqMXP68ER.KQtRy6._EHBVLci2QqZvdPiiIfe90LG5qoYtFHbbOAQD_6kXW3oI", path="/", domain=".dichtienghoa.com")
jar.set("_ga", "GA1.2.1746496857.1758984414", path="/", domain="dichtienghoa.com")
jar.set("_gid", "GA1.2.1310214375.1758984414", path="/", domain=".dichtienghoa.com")
jar.set("_gat", "1", path="/", domain=".dichtienghoa.com")
jar.set("_ga_VDQ6D99ZVH", "GS2.2.s1758984414`$o1`$g0`$t1758984414`$j60`$l0`$h0", path="/", domain=".dichtienghoa.com")

for chapter_content in chapters:
    # chapter_content = open("D:\\a.txt", "r", encoding='utf-8')
    # data for dichtienghoa.com
    data_dichtienghoa = {
        't': chapter_content,
        'tt': "vi"
    }
    
    #data for dichnhanh.com
    data_dichnhanh = {
        'type': "Ancient",
        'enable_analyze': "1",
        'enable_fanfic': "0",
        'mode': "vi",
        'text': chapter_content,
    }
    data_dichngay = {
        'content': chapter_content,
        'tl': "vi"
    }
    payload = json.dumps(data_dichngay)
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3',
        # 'Content-Type': 'application/json'
        "Content-Length": str(len(payload)),
    }
    try:
        r = requests.post(url=API_ENDPOINT, data=payload, headers=headers, timeout=30)

        if (r.status_code == 200):
            result = r.json()
            # append_to_docx("D:\\trieuhoansu.docx", result["data"]) # dich tieng hoa
            append_to_docx("D:\\trieuhoansu_dichngay.docx", result["data"]["content"]) # dich nhanh
        else:
            print(f"{Fore.RED}API dịch trả về lỗi. Mã lỗi: {r.status_code}")
        time.sleep(random.randint(2,5))
    except requests.exceptions.RequestException as e:
        print(f"{Fore.RED}Lỗi khi truy cập gọi api dịch: {e}")
